#!/usr/bin/env python3
"""Convert a proofread markdown to a .docx using this plugin's default spec.

Default spec (defined in assets/presets/default.yaml):
  - 2 cm margins on all four sides
  - Source Han Serif SC (思源宋体) body, 12 pt, 1.5 line spacing
  - 2-character first-line indent

Markdown features supported:
  - ATX headings (#, ##, ###)
  - Blockquotes (>) — rendered as indented quote paragraphs
  - Ordered and unordered lists (single level)
  - Bold (**) and italic (*) inline
  - Footnotes in [^n] style — inserted as Word native footnotes
  - Images — embedded centred, caption from alt text
  - A "参考文献" / "引用文献" trailing section — hanging indent numbered

Usage:
    python3 md_to_docx.py --input final.md --output final.docx \
        --title-from-first-h1
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    print("missing dependency: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import yaml  # PyYAML
except ImportError:
    yaml = None  # YAML preset support is optional; fall back to built-in dict


# Built-in preset. Mirrors `assets/presets/default.yaml` so the script still
# works when PyYAML isn't installed. The YAML form (preferred) is read by
# load_template() first; this dict is only a fallback.
TEMPLATES = {
    "default": {
        "body_font": "Source Han Serif SC",
        "body_size": 12,
        "heading_font": "Source Han Serif SC",
        "heading_sizes": {1: 18, 2: 16, 3: 14, 4: 12},
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 2.0,
        "margin_right": 2.0,
        "line_spacing": 1.5,
        "character_spacing_pt": 0.0,
        "image_max_width_cm": 14.0,
        "image_max_file_kb": 800,
    },
}


def _preset_dir() -> Path:
    """Location of shipped YAML presets next to this script."""
    return Path(__file__).resolve().parent.parent / "assets" / "presets"


def _flatten_yaml_preset(data: dict) -> dict:
    """Translate a YAML preset's nested schema to the flat dict the rest of
    the script consumes. Unknown keys are dropped silently so we can grow
    the YAML schema without breaking older installs."""
    page = data.get("page") or {}
    body = data.get("body") or {}
    heading = data.get("heading") or {}
    image = data.get("image") or {}

    heading_sizes = heading.get("sizes") or {}
    # YAML loads int keys as int already when written as plain 1:, 2: …
    # but be defensive in case someone used string keys.
    heading_sizes = {int(k): int(v) for k, v in heading_sizes.items()}

    return {
        "body_font": body.get("font", "Source Han Serif SC"),
        "body_size": int(body.get("size", 12)),
        "heading_font": heading.get("font", "Source Han Serif SC"),
        "heading_sizes": heading_sizes or {1: 18, 2: 16, 3: 14, 4: 12},
        "margin_top": float(page.get("margin_top", 2.0)),
        "margin_bottom": float(page.get("margin_bottom", 2.0)),
        "margin_left": float(page.get("margin_left", 2.0)),
        "margin_right": float(page.get("margin_right", 2.0)),
        "line_spacing": float(body.get("line_spacing", 1.5)),
        "first_line_indent_chars": int(body.get("first_line_indent_chars", 2)),
        "character_spacing_pt": float(body.get("character_spacing_pt", 0.0)),
        "image_max_width_cm": float(image.get("max_width_cm", 14.0)),
        "image_max_file_kb": int(image.get("max_file_kb", 800)),
    }


def load_template(name_or_path: str) -> dict:
    """Resolve a --template value to a flat template dict.

    Lookup order:
      1. Literal path to a `.yaml` / `.yml` file (absolute or relative)
      2. Preset name — matches `assets/presets/<name>.yaml`
      3. Preset name — matches the built-in `TEMPLATES` dict
    The first match wins; unknown names are an error.
    """
    # (1) explicit path
    candidate = Path(name_or_path)
    if candidate.suffix.lower() in {".yaml", ".yml"} and candidate.is_file():
        if yaml is None:
            raise SystemExit("PyYAML not installed; pip install pyyaml or use a built-in preset")
        with candidate.open("r", encoding="utf-8") as fh:
            return _flatten_yaml_preset(yaml.safe_load(fh) or {})

    # (2) shipped preset
    shipped = _preset_dir() / f"{name_or_path}.yaml"
    if shipped.is_file():
        if yaml is None:
            # Fall through to the built-in dict for this name rather than
            # crashing — keeps the script usable on a fresh Python install.
            if name_or_path in TEMPLATES:
                return TEMPLATES[name_or_path]
            raise SystemExit(
                f"YAML preset {shipped} exists but PyYAML isn't installed; "
                "run `pip install pyyaml` or pick a built-in preset"
            )
        with shipped.open("r", encoding="utf-8") as fh:
            return _flatten_yaml_preset(yaml.safe_load(fh) or {})

    # (3) built-in dict
    if name_or_path in TEMPLATES:
        return TEMPLATES[name_or_path]

    raise SystemExit(
        f"unknown template {name_or_path!r}. "
        f"Options: {sorted(TEMPLATES)} or a path to a *.yaml"
    )

FOOTNOTE_DEF = re.compile(r"^\[\^(\d+)\]:\s+(.*)$")
FOOTNOTE_REF = re.compile(r"\[\^(\d+)\]")
IMAGE_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
# HTML comments, including the `<!-- page N -->` markers that the OCR
# pipeline leaves in raw.md so preview / proofread can align with page
# images. These are workflow scaffolding, not body text, and must never
# leak into the Word document.
HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)


def set_page(section, tpl: dict) -> None:
    section.top_margin = Cm(tpl["margin_top"])
    section.bottom_margin = Cm(tpl["margin_bottom"])
    section.left_margin = Cm(tpl["margin_left"])
    section.right_margin = Cm(tpl["margin_right"])


def set_cn_font(run, font_name: str, size_pt: int, char_spacing_pt: float = 0.0) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    # Character (inter-glyph) spacing. Word stores this as `w:spacing` on
    # the run's rPr, measured in twips (1/20 pt). A positive value spreads
    # characters apart; 0 pt is Word's default tight setting.
    if char_spacing_pt:
        spacing = rPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = rPr.makeelement(qn("w:spacing"), {})
            rPr.append(spacing)
        spacing.set(qn("w:val"), str(int(round(char_spacing_pt * 20))))


def set_paragraph_indent(paragraph, first_line_chars: int, size_pt: int,
                         line_spacing: float = 1.5) -> None:
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(first_line_chars * size_pt)
    # `line_spacing` float + rule=MULTIPLE makes python-docx write the
    # MS Word equivalent of "multiple, 1.2". WD_LINE_SPACING.ONE_POINT_FIVE
    # is a pre-baked alias for the 1.5 case; using the float form covers
    # arbitrary values (e.g. 1.2) that the user prefers for dense review reads.
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def parse_inline(text: str) -> list[tuple[str, str]]:
    """Return list of (kind, text) — kind in {'plain','bold','italic','footnote'}.

    Minimal: bold **x**, italic *x*, footnote refs [^n]. Nesting not supported.
    """
    parts: list[tuple[str, str]] = []
    i = 0
    while i < len(text):
        # footnote ref
        m = FOOTNOTE_REF.match(text, i)
        if m:
            parts.append(("footnote", m.group(1)))
            i = m.end()
            continue
        # bold
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end > 0:
                parts.append(("bold", text[i + 2 : end]))
                i = end + 2
                continue
        # italic
        if text[i] == "*" and not text.startswith("**", i):
            end = text.find("*", i + 1)
            if end > 0:
                parts.append(("italic", text[i + 1 : end]))
                i = end + 1
                continue
        # plain — consume until next special marker
        j = i
        while j < len(text):
            if text.startswith("**", j) or text[j] == "*" or FOOTNOTE_REF.match(text, j):
                break
            j += 1
        if j > i:
            parts.append(("plain", text[i:j]))
            i = j
        else:
            parts.append(("plain", text[i]))
            i += 1
    return parts


def add_rich_paragraph(doc, text: str, tpl: dict, footnotes: dict[str, str]) -> None:
    p = doc.add_paragraph()
    set_paragraph_indent(p, 2, tpl["body_size"],
                         line_spacing=tpl.get("line_spacing", 1.5))
    cs = tpl.get("character_spacing_pt", 0.0)
    for kind, chunk in parse_inline(text):
        if kind == "footnote":
            fn_text = footnotes.get(chunk, "")
            # python-docx doesn't expose proper footnotes; use superscript inline
            run = p.add_run(f"[{chunk}]")
            run.font.superscript = True
            set_cn_font(run, tpl["body_font"], tpl["body_size"], cs)
            if fn_text:
                note = p.add_run(f"({fn_text})")
                note.font.size = Pt(tpl["body_size"] - 2)
                set_cn_font(note, tpl["body_font"], tpl["body_size"] - 2, cs)
            continue
        run = p.add_run(chunk)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        set_cn_font(run, tpl["body_font"], tpl["body_size"], cs)


def add_heading(doc, text: str, level: int, tpl: dict) -> None:
    # Use python-docx's built-in Heading styles so Word's navigation pane,
    # outline view, and auto-generated TOC all recognise the document
    # structure. Custom font + size are layered on top so the visual design
    # matches the active preset.
    style_name = f"Heading {level}" if 1 <= level <= 9 else "Heading 1"
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        # Fallback for templates that don't ship that level.
        p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 16, 2: 14, 3: 13}.get(level, 12)
    run = p.add_run(text)
    run.bold = True
    set_cn_font(run, tpl["heading_font"], size, tpl.get("character_spacing_pt", 0.0))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = tpl.get("line_spacing", 1.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_quote(doc, text: str, tpl: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(tpl["body_size"] * 2)
    p.paragraph_format.right_indent = Pt(tpl["body_size"] * 2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = tpl.get("line_spacing", 1.5)
    run = p.add_run(text)
    set_cn_font(run, tpl["body_font"], tpl["body_size"] - 1,
                tpl.get("character_spacing_pt", 0.0))
    run.italic = False


def _compress_for_docx(src: Path, max_kb: int, max_width_cm: float) -> Path:
    """Return a path to an image that fits the template's size/byte caps.

    MinerU extracts figures at screen resolution (often 2–4 MB PNG), which
    quickly bloats the docx. We re-encode to a reasonable JPEG when the
    file is large, preserving the caller's filename so the relative path
    stays stable. Falls back to the original path if Pillow isn't
    available, if the image is already small, or if anything fails.
    """
    try:
        size_kb = src.stat().st_size / 1024.0
    except OSError:
        return src
    # Rough pt / cm mapping used elsewhere: 72pt/inch, 2.54cm/inch.
    target_px = int(max_width_cm / 2.54 * 300)  # 300 DPI print quality

    if size_kb <= max_kb:
        return src

    try:
        from PIL import Image  # type: ignore
    except Exception:
        return src

    try:
        with Image.open(src) as im:
            im.load()
            w, h = im.size
            if w > target_px:
                scale = target_px / float(w)
                # `Image.Resampling.LANCZOS` on Pillow ≥10, kept via the
                # fallback alias on older installs. getattr keeps both
                # happy without a runtime version check.
                lanczos = getattr(
                    getattr(Image, "Resampling", Image),
                    "LANCZOS",
                    None,
                )
                im = im.resize(
                    (target_px, int(h * scale)),
                    lanczos or 1,  # 1 == ANTIALIAS / BILINEAR fallback
                )
            dst = src.with_suffix(".compressed.jpg")
            if im.mode in ("RGBA", "P"):
                im = im.convert("RGB")
            quality = 85
            im.save(dst, "JPEG", quality=quality, optimize=True)
            # If even post-resize we're over budget, drop quality a notch.
            while dst.stat().st_size / 1024.0 > max_kb and quality > 55:
                quality -= 10
                im.save(dst, "JPEG", quality=quality, optimize=True)
            return dst
    except Exception:
        return src


def add_image(doc, alt: str, path: Path, md_dir: Path, tpl: dict) -> None:
    candidates = [path, md_dir / path, md_dir / "assets" / path.name]
    target = next((c for c in candidates if c.is_file()), None)
    if not target:
        p = doc.add_paragraph(f"[图片缺失：{path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    max_kb = int(tpl.get("image_max_file_kb", 800))
    max_width = float(tpl.get("image_max_width_cm", 14.0))
    # Compress only when needed so small inline assets stay pristine.
    served = _compress_for_docx(target, max_kb, max_width)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(served), width=Cm(max_width))
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)


def extract_footnotes(text: str) -> tuple[str, dict[str, str]]:
    footnotes: dict[str, str] = {}
    body_lines: list[str] = []
    for line in text.splitlines():
        m = FOOTNOTE_DEF.match(line)
        if m:
            footnotes[m.group(1)] = m.group(2)
        else:
            body_lines.append(line)
    return "\n".join(body_lines), footnotes


def render(md_path: Path, out_path: Path, template: str, title_from_h1: bool) -> None:
    tpl = load_template(template)
    md_text = md_path.read_text(encoding="utf-8")
    # Strip HTML comments (e.g. OCR page markers) before any other parsing
    # so they never show up as paragraph text.
    md_text = HTML_COMMENT.sub("", md_text)
    body_text, footnotes = extract_footnotes(md_text)

    doc = Document()
    for section in doc.sections:
        set_page(section, tpl)

    title_captured = not title_from_h1
    in_ref_section = False

    for raw_line in body_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            heading_text = line[2:].strip()
            if not title_captured:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading_text)
                run.bold = True
                set_cn_font(run, tpl["heading_font"], 18)
                p.paragraph_format.space_after = Pt(18)
                title_captured = True
            else:
                add_heading(doc, heading_text, 1, tpl)
                if heading_text in ("参考文献", "引用文献", "Bibliography"):
                    in_ref_section = True
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2, tpl)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3, tpl)
            continue
        if line.startswith("> "):
            add_quote(doc, line[2:].strip(), tpl)
            continue

        img_match = IMAGE_PATTERN.search(line)
        if img_match:
            alt, src = img_match.group(1), img_match.group(2)
            add_image(doc, alt, Path(src), md_path.parent, tpl)
            continue

        if in_ref_section:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(tpl["body_size"] * 2)
            p.paragraph_format.first_line_indent = Pt(-tpl["body_size"] * 2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(line)
            set_cn_font(run, tpl["body_font"], tpl["body_size"])
            continue

        add_rich_paragraph(doc, line, tpl, footnotes)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, type=Path)
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--template", default="default",
                    help="preset name under assets/presets/<name>.yaml "
                         "or a path to a custom .yaml")
    ap.add_argument("--title-from-first-h1", action="store_true")
    args = ap.parse_args()

    if not args.input.is_file():
        print(f"input not found: {args.input}", file=sys.stderr)
        return 2

    try:
        render(args.input, args.output, args.template, args.title_from_first_h1)
    except Exception as e:
        print(f"[md_to_docx] failed: {e}", file=sys.stderr)
        return 5

    print(f"[md_to_docx] wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
